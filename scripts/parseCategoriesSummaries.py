# python3 parseSummaryKeywords.py categories.txt summaries.txt InfoFolder
# output: InfoFolder/categories.js 
#		  InfoFolder/summaries.js

# categories.js = {
#     "1": {
#         "name": "Yugtun Qaneryaraq -- Yup'ik Language"
#         "children": 0,
#         "videoNumbers": [ videoNumbers ]
#     };

# summaries.js = {
#     "1": {
#         "videoID": "cpb-aacip-127-009w0z0q.h264",
#         "tags": [ categoryNumbers ],
#         "english": {
#             "summary": [ "Joshua Phillip (Ac'urunaq; nicknamed Maqista)", ],
#             "timestamps": [ "00:00 - Yup'ik langauge", ]
#         },
#         "yugtun": {
#             "summary": [ "Ac'urunaq (nickname-aqluku Maqista)" ],
#             "timestamps": [ "00:00 - Yugtun qalarcaram tamanritlerkaa", ]
#         }
#     };

# import sys
import os
from collections import defaultdict, OrderedDict
import re
import json
# import pdb
import string
from lxml import etree
from pathlib import Path
import glob
import docx
import csv

videoNumAndDuration = OrderedDict([
("cpb-aacip-127-009w0z0q.h264",("1","1:02:35")),#"1:02:35.230000",
("cpb-aacip-127-00ns1t6z.h264",("2","31:22")),#"0:31:22.548333",
("cpb-aacip-127-010p2r15.h264",("3","22:27")),#"0:22:27.468333",
("cpb-aacip-127-032282mq.h264",("4","12:53")),#"0:12:53.313333",
("cpb-aacip-127-03cz8zdq.h264",("5","22:23")),#"0:22:23.665000",
("cpb-aacip-127-0644j324.h264",("6","21:05")),#"0:21:05.690000",
("cpb-aacip-127-0644j37r.h264",("7","10:06")),#"0:10:06.385000",
("cpb-aacip-127-06g1jzz6.h264",("8","21:59")),#"0:21:59.275000",
("cpb-aacip-127-06g1k008.h264",("9","22:42")),#"0:22:42.216667",
("cpb-aacip-127-085hqfqk.h264",("10","31:54")),#"0:31:54.616667",
("cpb-aacip-127-09w0vx3c.h264",("11","18:02")),#"0:18:02.345000",
("cpb-aacip-127-10jsxpvr.h264",("12","22:23")),#"0:22:23.063333",
("cpb-aacip-127-10jsxpwg.h264",("13","22:27")),#"0:22:27.801667",
("cpb-aacip-127-10jsxpx6.h264",("14","21:49")),#"0:21:49.231667",
("cpb-aacip-127-149p8hcz.h264",("15","22:02")),#"0:22:02.343333",
("cpb-aacip-127-14nk9d19.h264",("16","21:55")),#"0:21:55.638333",
("cpb-aacip-127-15p8d31m.h264",("17","21:35")),#"0:21:35.718333",
("cpb-aacip-127-16pzgr3f.h264",("18","21:28")),#"0:21:28.411667",
("cpb-aacip-127-18rbp380.h264",("19","22:30")),#"0:22:30.705000",
("cpb-aacip-127-20fttjr7.h264",("20","22:33")),#"0:22:33.073333",
("cpb-aacip-127-225b014v.h264",("21","22:05")),#"0:22:05.046667",
("cpb-aacip-127-23612qck.h264",("22","1:01:44")),#"1:01:44.646667",
("cpb-aacip-127-257d81kk.h264",("23","18:22")),#"0:18:22.031667",
("cpb-aacip-127-257d81m9.h264",("24","22:06")),#"0:22:06.348333",
("cpb-aacip-127-25k98x78.h264",("25","21:23")),#"0:21:23.206667",
("cpb-aacip-127-25x69tk3.h264",("26","22:26")),#"0:22:26.066667",
("cpb-aacip-127-26xwdh7k.h264",("27","21:17")),#"0:21:17.801667",
("cpb-aacip-127-27zkh5wz.h264",("28","13:50")),#"0:13:50.801667",
("cpb-aacip-127-28ncjzpp.h264",("29","50:28")),#"0:50:28.726667",
("cpb-aacip-127-2908kvc9.h264",("30","22:12")),#"0:22:12.888333",
("cpb-aacip-127-322bvwfr.h264",("31","20:47")),#"0:20:47.071667",
("cpb-aacip-127-33rv1b2m.h264",("32","10:11")),#"0:10:11.656667",
("cpb-aacip-127-34fn33pp.h264",("33","22:23")),#"0:22:23.465000",
("cpb-aacip-127-34sj40gx.h264",("34","15:34")),#"0:15:34.870000",
("cpb-aacip-127-35gb5sdj.h264",("35","22:23")),#"0:22:23.698333",
("cpb-aacip-127-35t76q2k.h264",("36","15:52")),#"0:15:52.086667",
("cpb-aacip-127-37hqc4tk.h264",("37","21:15")),#"0:21:15.698333",
("cpb-aacip-127-37vmd1rq.h264",("38","10:14")),#"0:10:14.960000",
("cpb-aacip-127-386hdxkk.h264",("39","21:17")),#"0:21:17.801667",
("cpb-aacip-127-386hdz1p.h264",("40","1:01:44")),#"1:01:44.946667",
("cpb-aacip-127-38w9gpk2.h264",("41","18:55")),#"0:18:55.663333",
("cpb-aacip-127-41mgqv1k.h264",("42","22:01")),#"0:22:01.743333",
("cpb-aacip-127-42n5thvk.h264",("43","20:13")),#"0:20:13.171667",
("cpb-aacip-127-440rz436.h264",("44","31:25")),#"0:31:25.588333",
("cpb-aacip-127-46254fdx.h264",("45","21:04")),#"0:21:04.888333",
("cpb-aacip-127-47rn8x0g.h264",("46","20:34")),#"0:20:34.158333",
("cpb-aacip-127-48ffbpfx.h264",("47","1:01:42")),#"1:01:42.045000",
("cpb-aacip-127-52j6qcr2.h264",("48","1:03:36")),#"1:03:36.855000",
("cpb-aacip-127-53wstz0s.h264",("49","22:25")),#"0:22:25.066667",
("cpb-aacip-127-54kkwqvx.h264",("50","22:07")),#"0:22:07.148333",
("cpb-aacip-127-54kkwrd6.h264",("51","1:01:16")),#"1:01:16.120000",
("cpb-aacip-127-558czhn9.h264",("52","10:33")),#"0:10:33.278333",
("cpb-aacip-127-57np5s39.h264",("53","18:40")),#"0:18:40.248333",
("cpb-aacip-127-59c5b6f6.h264",("54","19:05")),#"0:19:05.306667",
("cpb-aacip-127-59c5b6gx.h264",("55","21:44")),#"0:21:44.626667",
("cpb-aacip-127-59q2c3zn.h264",("56","15:16")),#"0:15:16.051667",
("cpb-aacip-127-60qrfsqn.h264",("57","22:06")),#"0:22:06.548333",
("cpb-aacip-127-612ngp56.h264",("58","15:38")),#"0:15:38.073333",
("cpb-aacip-127-623bkbxr.h264",("59","8:26")),#"0:08:26.888333",
("cpb-aacip-127-6341p189.h264",("60","9:44")),#"0:09:44.665000",
("cpb-aacip-127-63fxpxg0.h264",("61","31:26")),#"0:31:26.290000",
("cpb-aacip-127-65h9w88k.h264",("62","1:00:49")),#"1:00:49.526667",
("cpb-aacip-127-65v6x4wh.h264",("63","19:48")),#"0:19:48.415000",
("cpb-aacip-127-66vx0v02.h264",("64","21:52")),#"0:21:52.235000",
("cpb-aacip-127-67wm3hhj.h264",("65","31:05")),#"0:31:05.201667",
("cpb-aacip-127-68kd59ft.h264",("66","22:22")),#"0:22:22.096667",
("cpb-aacip-127-69m37zr9.h264",("67","26:53")),#"0:26:53.658333",
("cpb-aacip-127-719kdf8m.h264",("68","10:37")),#"0:10:37.715000",
("cpb-aacip-127-719kdfjd.h264",("69","1:02:28")),#"1:02:28.088333",
("cpb-aacip-127-74cnpfqb.h264",("70","13:38")),#"0:13:38.790000",
("cpb-aacip-127-74cnpfr2.h264",("71","20:37")),#"0:20:37.761667",
("cpb-aacip-127-74qjqbsc.h264",("72","21:34")),#"0:21:34.516667",
("cpb-aacip-127-752fr7q2.h264",("73","58:49")),#"0:58:49.878333",
("cpb-aacip-127-7634tx89.h264",("74","22:16")),#"0:22:16.358333",
("cpb-aacip-127-76f1vspw.h264",("75","16:55")),#"0:16:55.246667",
("cpb-aacip-127-773txm5v.h264",("76","22:19")),#"0:22:19.761667",
("cpb-aacip-127-784j17vw.h264",("77","40:01")),#"0:40:01.286667",
("cpb-aacip-127-79h44thw.h264",("78","22:21")),#"0:22:21.161667",
("cpb-aacip-127-805x6mnx.h264",("79","18:30")),#"0:18:30.940000",
("cpb-aacip-127-816m99hb.h264",("80","1:02:05")),#"1:02:05.166667",
("cpb-aacip-127-81jhb5w6.h264",("81","22:26")),#"0:22:26.868333",
("cpb-aacip-127-84mkm6tr.h264",("82","6:20")),#"0:06:20.166667",
("cpb-aacip-127-85n8pw5n.h264",("83","22:27")),#"0:22:27.168333",
("cpb-aacip-127-86nzsk0r.h264",("84","13:49")),#"0:13:49.300000",
("cpb-aacip-127-87brvbt5.h264",("85","14:30")),#"0:14:30.606667",
("cpb-aacip-127-87pnw7r9.h264",("86","17:36")),#"0:17:36.420000",
("cpb-aacip-127-881jx4d6.h264",("87","51:26")),#"0:51:26.883333",
("cpb-aacip-127-89280sv6.h264",("88","1:04:11")),#"1:04:11.823333",
("cpb-aacip-127-89r22kjv.h264",("89","21:45")),#"0:21:45.761667",
("cpb-aacip-127-90rr586b.h264",("90","26:40")),#"0:26:40.946667",
("cpb-aacip-127-91sf7xm0.h264",("91","29:55")),#"0:29:55.935000",
("cpb-aacip-127-945qg6jq.h264",("92","22:07")),#"0:22:07.448333",
("cpb-aacip-127-956djvkd.h264",("93","22:10")),#"0:22:10.951667",
("cpb-aacip-127-96wwq9xv.h264",("94","22:22")),#"0:22:22.296667",
("cpb-aacip-127-988gtvzp.h264",("95","3:28")),#"0:03:28.166667",
("cpb-aacip-127-98mcvs0p.h264",("96","22:23")),#"0:22:23.165000",
])

categoryReplacements = {
"\tTegganeq -- Elder":"Tegganeq -- Elder",
"-- Qemagciyaraq -- Storage":"Qemagciyaraq -- Storage",
"Aipanglleq -- Getting A Spouse":"Aipangyaraq -- Getting A Spouse",
"Aipaqellriik, Aipaqelriik, Nulirqellriik -- Married Couple":"Aipangyaraq, Aipaqsaraq -- Marriage",
"Aklut -- Clothes":"Akluq, Aklut -- Clothing, Possessions",
"Akusrarun --  Mischief, Misconduct":"Akusrarun -- Mischief, Misconduct",
"Alangru -- Ghost":"Alangru, Aliurtuq -- Ghost, Haunting",
"Amiq, Qecik -- Skin":"Amiq, Qecik -- Skin, Hide",
"Aviukaqsaraq -- Food And Water Offering":"Aviukaryaraq -- Food And Water Offering",
"Calricaraq -- Health and Wellness":"Calricaraq -- Health & Wellness",
"Calrunani Unguvamaniartuq -- To Live and Long and Healthy Life":"Calrunani Unguvamaniartuq -- To Live a Long and Healthy Life",
"Canek, Can'get, Caranglluk -- Grass":"Canek, Can'get, Caranglluk, Evek -- Grass",
"Ella Qupsaraa -- The World, Universe Cracks":"Ellam Qupsaraa -- The World, Universe Cracks",
"Ellalliurcuutet -- Rain Gear":"Ellaliurcuutet -- Rain Gear",
"Ellam Yua -- Spirit Of The Universe":"Ellam Yua, Cellam Yua, Cillam Cua -- Spirit Of The Universe",
"Ellam Yua -- Spirit of the Universe":"Ellam Yua, Cellam Yua, Cillam Cua -- Spirit Of The Universe",
"Ellam Yua -- The Spirit of the Universe":"Ellam Yua, Cellam Yua, Cillam Cua -- Spirit Of The Universe",
"Ellam Yua, Cellam Yua, Cillam Cua -- Spirit of the Universe":"Ellam Yua, Cellam Yua, Cillam Cua -- Spirit Of The Universe",
"Enepiaq, Nepiaq -- Sod-House":"Nepiaq, Enepiaq, Enpiaq -- Sod-house",
"Imarpigmiutaat, Unkumiutaat -- Marine Animals":"Imarpigmiutaat, Unkumiutaat, Mermiutaat -- Marine Animals",
"Imarpigmiutaat, Unkumkiutaat, Mermiutaat -- Marine Animals":"Imarpigmiutaat, Unkumiutaat, Mermiutaat -- Marine Animals",
"Iraluq -- Month":"Iraluq -- Moon, Month",
"Ircenrraat -- Legendary Little People":"Ircenrraat -- Legendary Little People [Mischievous]",
"Kaignaq, Piitnaq -- Time of Famine, Need":"Kaignaq, Piitnaq -- Time of Hunger, Need",
"Kapkaanaq -- Steel Trap":"Kapkaanaq -- Conibear, Steel Trap",
"Kelugkaq -- Grass Grass for Weaving Mats":"Kelugkaq -- Course Grass for Weaving Mats",
"Kit'elleq -- Falling into Water":"Kit'elleq -- Falling Into Water",
"Kumlivik Freezer":"Kumlivik -- Freezer",
"Maligtaqustiin Yuunertucia -- Those Who Heed Advice Live Long":"Maligtaqustiin Yuunertutucia -- Those Who Heed Advice Live Long",
"Maligtaquyaraq -- Obedience":"Maligtaquyaraq, Niisngayaraq -- Obedience",
"Maqi -- Steambath":"Maqi -- Steambath, Fire Bath",
"Meluk -- Roe":"Meluk, Imlauk -- Roe",
"Naulluun, Qenan -- Sickness":"Naulluun, Qenan, Nangtequn -- Sickness",
"Naumriit Enenguat -- Households Made of Plants":"Naumalriit Enenguat -- Households Made of Plants",
"Naunrat -- Plants":"Naunraat -- Plants",
"Neqpik -- Fish":"Neqa, Neqet -- Fish",
"Nerangnaqsaraq -- Subsistence":"Nerangnaqsaraq, Yuungnaqsaraq -- Subsistence",
"Pellalleq -- Getting Lost":"Pellaalleq -- Getting Lost",
"Pissuutet, Cassuutet, Piliat, Saskut, Ayagassuutet -- Tools, Crafts, Weapon":"Pissuutet, Cassuutet, Piliat, Saskut, Ayagassuutet -- Tools, Crafts, Weapons",
"Piunrillret, Yuunrillret, Tuqullret -- Afterlife":"Piunrillret, Yuunrillret, Tuqullret, Catairutellret -- Afterlife",
"Piunrissiyaayuitellrat Naulluussiyaayuunateng-llu -- Fewer Sickness and Death":"Piunrissiyaayuitellrat Naulluussiyaayuunateng-llu -- Fewer Death and Sickness",
"Qanruyutet -- Traditional Wisdom, Wise Words":"Qanruyutet, Qaneryarat -- Traditional Wisdom, Wise Words",
"Qasgiq, Qaygiq -- Community, Men's House":"Qasgiq, Qaygiq -- Community or Men's House",
"Tuntu -- Reindeer":"Tuntu -- Reindeer, Caribou",
"Tuqu Naulluun-llu, Qenan-llu -- Death and Sickness":"Tuqu Naulluun-llu, Qenan-llu -- Death & Sickness",
"Tutmarngaingurmun Ciqicarat -- Disposal Where People Won't Trod":"Tutmarngailngurmun Ciqicarait -- Disposal Where People Won't Trod",
"Tuunriyaraq, Angalkiryaraq -- Shamanistic Acts, Rituals":"Tuunriyaraq, Angalkiyaraq -- Shamanistic Acts, Rituals",
"Unangkengaita Auluksarait -- Taking Care of Catch":"Ungangkengaita Auluksarait -- Taking Care of Catch",
"Unangkengaita Auluksarat -- Taking Care of Catch":"Ungangkengaita Auluksarait -- Taking Care of Catch",
"Uqren Aarnarqucia -- The Danger of Leeward Sides":"Uqrem Aarnarqucia -- The Danger of Leeward Sides",
"Yaqulek, Yaqulget -- Birds":"Yaqulek, Yaqulget, Tengmiaq, Tengmiat -- Birds",
"Yaqulek, Yaqulget, Tengmiat, Tengmiat -- Birds":"Yaqulek, Yaqulget, Tengmiaq, Tengmiat -- Birds",
"Yuungcarviit -- Clinics and Hospitals":"Yuungcarviit -- Clinics & Hospitals",
"Yuuyaraq -- Way of Life":"Yuuyaraq -- Way Of Life",
"Aglenraraq -- First Year Of Menstruation":"Aglenraraq -- First Year of Menstruation",
"Naunraat Atsat-llu -- Plants and Berries":"Naunraat Atsat-llu -- Plants And Berries",
"-- Taangaq -- Alcohol":"Taangaq -- Alcohol",
"Aipangyaraq -- Getting a Spouse":"Aipangyaraq -- Getting A Spouse",
"Anglicarillerkaq -- Child Rearing":"Anglicarillerkaq, Tukercaryaraq -- Child Rearing",
"Caniqerrilin, Caniqerrilitaq -- Woven Grass Blet Blocking Door During Instruction":"Caniqerrilin, Caniqerrilitaq -- Woven Grass Belt Blocking Door During Instruction",
"Nuliangyaraq, Nulirturyaraq -- How to Work With a Wife":"Nuliangyaraq, Nulirturyaraq -- How To Work With A Wife",
"Qengaruk -- Shelter":"Qengaruk -- Snowbank",
"Taangaq --  Alcohol":"Taangaq -- Alcohol",
"Tuntuciuryaraq -- Reindeer Herding":"Tunciuryaraq -- Reindeer Herding",
"Uingyaraq -- How to Work With a Husband":"Uingyaraq -- How To Work With A Husband",
"Yugtun Qalarcaraq -- Yup'ik Language":"Yugtun Qalarcaraq, Qaneryaraq -- Yup'ik Language",
}

rx_dict = {
	'videoID' : re.compile(r'# (?P<videoNumber>\d+) (?P<videoID>\S+)'),
	'lang' : re.compile(r'## (?P<lang>Yugtun|English)'),
	'title' : re.compile(r'## Title'),
	'date' : re.compile(r'## Date'),
	'tags' : re.compile(r'## Keywords'),
	'summary' : re.compile(r'### Summary'),
	'timestamps' : re.compile(r'### Timestamps')
}

def _parse_line(line):
	for key, rx in rx_dict.items():
		match = rx.search(line)
		if match:
			return key, match
	return None, None

def parse_summaries(filepath):
	data = defaultdict(dict)
	with open(filepath, 'r') as file_object:
		line = file_object.readline()
		while line:
			key, match = _parse_line(line)
			# print(f'{key} {line}')

			if key == None:
				line = file_object.readline()

			if key == 'videoID':
				videoNumber = match.group('videoNumber')
				videoID = match.group('videoID')
				data[videoNumber]['videoID'] = videoID
				line = file_object.readline()

				title = "Title Placeholder"
				date = "Unknown"
				data[videoNumber]['title'] = title
				data[videoNumber]['date'] = date

			if key == 'title':
				line = file_object.readline()
				while True:
					if line.strip():
						if line.strip()[0] == '#':
							break
						title = line.strip()
					line = file_object.readline()
				data[videoNumber]['title'] = title

			if key == 'date':
				line = file_object.readline()
				while True:
					if line.strip():
						if line.strip()[0] == '#':
							break
						date = line.strip()
					line = file_object.readline()
				data[videoNumber]['date'] = date

			if key == 'tags':
				tags = []
				line = file_object.readline()
				while True:
					if line.strip():
						if line.strip()[0] == '#':
							break
						tags.append(line.strip())
					line = file_object.readline()
				data[videoNumber]['tags'] = tags

			if key == 'lang':
				lang = match.group('lang').lower()
				line = file_object.readline()

			if key == 'summary':
				summary_text = []
				line = file_object.readline()
				while True:
					if line.strip():
						if line.strip()[0] == '#':
							break
						summary_text.append(line.strip())
					line = file_object.readline()

			if key == 'timestamps':
				timestamps = []		
				line = file_object.readline()
				while True:
					if line.strip():
						if line.strip()[0] == '#':
							break
						timestamps.append(line.strip())
					line = file_object.readline()
				data[videoNumber][lang] = {'summary':summary_text,
										   'timestamps':timestamps}
	restrucutredData = {}
	for i in data:
		aapb = data[i]["videoID"]
		restrucutredData[aapb] = {}
		restrucutredData[aapb]["videoID"] = aapb
		restrucutredData[aapb]["title"] = data[i]["title"]
		restrucutredData[aapb]["time"] = videoNumAndDuration[aapb][1]
		restrucutredData[aapb]["metadata"] = {
			"Date":data[i]["date"],
			"Short Summary:":[data[i]["yugtun"]['summary'],data[i]["english"]['summary']]}
		new_timestamps = []
		for t, timestamp in enumerate(data[i]["yugtun"]["timestamps"]):
			time, yugtun = timestamp.split(" - ",1)
			english = data[i]["english"]["timestamps"][t].split(" - ",1)[1]
			new_timestamps.append([time,yugtun,english])
		restrucutredData[aapb]["summary"] = new_timestamps
		restrucutredData[aapb]["elderTags"] = []
		restrucutredData[aapb]["tags"] = data[i]["tags"]

	# return data
	return restrucutredData

coreykeywords = [
"AAPB ID:",
"Title:",
"Date:",
"Date (if known):",
"Short Summary:",
"Summary:",
"Genre 1:",
"Genre 2:",
"Quality:",
"Genre 4 (Location):",
"Location:",
"Name 1:",
"Role 1:",
"Name 2:",
"Role 2:",
"TAGS:",
]

def parseCoreySummaries(files):
	data = defaultdict(dict)

	def tagFixing(w):
		w = w.replace("’","'")
		w = w.replace(" - "," -- ")
		w = w.replace(" – "," -- ")
		return w

	# files_one = ['../data/WoW-Corey/cpb-aacip-127-76f1vspw.h264 - Summary.docx',]
	# for file in files_one:
	for file in files:
		print(file)
		doc = docx.Document(file)
		aapb = ''
		videodata = {
			"videoID":"",
			"title":"Title Placeholder",
			"time":"",
			"metadata":{
				"Date":"Unknown",
			},
			"summary":[],
			"elderTags":[],
			"tags":[]
		}

		i = 0
		while i < len(doc.paragraphs):
			line = doc.paragraphs[i].text.strip()
			# print(line)
			# if not line:
			# 	i += 1
			# 	continue
			
			if line == "AAPB ID:":
				i += 1
				if i >= len(doc.paragraphs):
					break
				aapb = doc.paragraphs[i].text
				if aapb in coreykeywords:
					print('no aapb id')
					break
				videodata['videoID'] = aapb
				videodata['time'] = videoNumAndDuration[aapb][1]
				i += 1
				continue

			elif line == "Title:":
				i += 1
				if i >= len(doc.paragraphs):
					break
				line = doc.paragraphs[i].text
				if line in coreykeywords:
					# print('no title for {aapb}')
					continue
				videodata['title'] = line
				i += 1
				continue

			elif line == "Summary:":
				summary = []
				i += 1
				if i >= len(doc.paragraphs):
					break
				while True:
					line = doc.paragraphs[i].text.strip()
					if line in coreykeywords:
						break
					# print(line)
					timestamp = doc.paragraphs[i].text.strip()
					english = doc.paragraphs[i+1].text.strip()
					yugtun = doc.paragraphs[i+2].text.strip()
					timestamp = re.sub(r'^0([1-9]):',r'\1:',timestamp) # remove excess numbering
					timestamp = re.sub(r'^00:','',timestamp) # remove excess numbering
					timestamp = re.sub(r'\..*','',timestamp) # remove excess numbering
					timestamp = timestamp.strip()
					summary.append([timestamp,yugtun,english])
					i += 3
					if i >= len(doc.paragraphs):
						break
				videodata['summary'] = summary

			elif line.startswith("Date") or \
				 line.startswith('Short Summary') or \
				 line.startswith('Genre') or \
				 line.startswith('Quality') or \
				 line.startswith('Location') or \
				 line.startswith('Name') or \
				 line.startswith('Role'):
				datatype = line
				datalines = []
				i += 1
				if i >= len(doc.paragraphs):
					break
				while True:
					line = doc.paragraphs[i].text
					if line in coreykeywords:
						break
					# print(line)
					datalines.append(line)
					i += 1
					if i >= len(doc.paragraphs):
						break
				if datatype.startswith("Date"):
					datatype = "Date"
				videodata['metadata'][datatype] = datalines

			elif line == "TAGS:":
				tags = []
				i += 1
				if i >= len(doc.paragraphs):
					break
				while True:
					line = doc.paragraphs[i].text
					if line in coreykeywords:
						break
					# print(line)
					tags.append(tagFixing(line))
					i += 1
					if i >= len(doc.paragraphs):
						break
				videodata['tags'] = tags
			
			else:
				print(line)
				i += 1

		leafTags = []
		level = 0
		for t,tag in enumerate(videodata['tags']): # only grab leaf tags
			level = tag.count('\t')
			if t == len(videodata['tags'])-1: # last tag
				leafTags.append(tag)
			elif videodata['tags'][t+1].count('\t') == level + 1:
				continue
			else:
				leafTags.append(tag)
		leafTags = [re.sub(r'^\t*[\d.]+ ','',t) for t in leafTags] # remove tabs and numbers
		videodata['tags'] = leafTags

		data[aapb] = videodata

	return data
			


def categoriesXML(filename, elderNames):
	tempfilename = '../data/categories-temp.txt'
	stack = []
	parse_line = re.compile(r'(?P<tabs>\s*)(?P<title>.*)')

	categories = etree.Element("categories")
	subcategory = None
	stack.append(categories)

	# add elderNames
	elders = sorted(list(elderNames))
	data = open(filename, "r").read()
	data = data.replace("INSERTELDERSHERE",'\n\t'.join(elders))
	with open(tempfilename, 'w') as out:
		out.write(data)

	with open(tempfilename, 'r') as file:
		for line in file:
			if not line.strip():
				continue
			level = len(stack) - 1

			match = parse_line.search(line)
			tabCount = len(match.group('tabs'))
			title = match.group('title')

			if tabCount > level:
				stack.append(subcategory)
				subcategory = etree.SubElement(stack[-1], 'subcategory', name=title)

			elif tabCount < level:
				stack = stack[:tabCount+1]
				subcategory = etree.SubElement(stack[-1], 'subcategory', name=title)

			else:
				subcategory = etree.SubElement(stack[-1], 'subcategory', name=title)

	# print(etree.tostring(categories, pretty_print=True, encoding='unicode'))
	return categories

def categoriesDictFunc(categoryMasterTree):
	categories = defaultdict(dict)

	# 4 layer deep subcategories, add another layer if more
	for i_0, subcategory_i in enumerate(categoryMasterTree):
		i = i_0 + 1
		i_index = str(i)
		categories[i_index]['name'] = subcategory_i.attrib['name']
		categories[i_index]['url'] = re.sub(r', | & | ~ | ','-', subcategory_i.attrib['name'].split('--')[0].strip().replace("'",""))
		categories[i_index]['images'] = []
		categories[i_index]['children'] = len(subcategory_i)
		categories[i_index]['videoNumbers'] = []
		for j_0, subcategory_j in enumerate(subcategory_i):
			j = j_0 + 1
			j_index = '.'.join([i_index,str(j)])
			categories[j_index]['name'] = subcategory_j.attrib['name']
			categories[j_index]['url'] = re.sub(r', | & | ~ | ','-', subcategory_j.attrib['name'].split('--')[0].strip().replace("'",""))
			categories[j_index]['images'] = []
			categories[j_index]['children'] = len(subcategory_j)
			categories[j_index]['videoNumbers'] = []
			for k_0, subcategory_k in enumerate(subcategory_j):
				k = k_0 + 1
				k_index = '.'.join([j_index,str(k)])
				categories[k_index]['name'] = subcategory_k.attrib['name']
				categories[k_index]['url'] = re.sub(r', | & | ~ | ','-', subcategory_k.attrib['name'].split('--')[0].strip().replace("'",""))
				categories[k_index]['images'] = []
				categories[k_index]['children'] = len(subcategory_k)
				categories[k_index]['videoNumbers'] = []
				for l_0, subcategory_l in enumerate(subcategory_k):
					l = l_0 + 1
					l_index = '.'.join([k_index,str(l)])
					categories[l_index]['name'] = subcategory_l.attrib['name']
					categories[l_index]['url'] = re.sub(r', | & | ~ | ','-', subcategory_l.attrib['name'].split('--')[0].strip().replace("'",""))
					categories[l_index]['images'] = []
					categories[l_index]['children'] = len(subcategory_l)
					categories[l_index]['videoNumbers'] = []

	return categories


def addElderIdentification(elderIdentifierFilename, summariesDict):
	elderNames = set()
	elderCat2Images = defaultdict(list)

	with open(elderIdentifierFilename, mode='r') as file:
		# csvFile = csv.reader(file, delimiter='\t')
		reader = csv.DictReader(file, delimiter='\t')
		
		unknownElderNumber = 1
		for lineDict in reader:
			#{'Photo of Elder': '', 
			#'Video #': 'cpb-aacip-127-00ns1t6z.h264.mov', 
			#'Name ': 'Ackiar', 
			#'English Name': 'Nick Lupie', 
			#'Village': 'Tuntutuliak', 
			#'Notes': '',
			#'Image': ''
			#print(lineDict)
			aapb = lineDict['Video #'].replace('.mov','')
			if aapb not in summariesDict:
				print(f'{aapb} - elder identification - missing aapb summary')
				continue

			elderString = lineDict['Name'].strip()
			unnamedWithVillage = False
			if elderString and lineDict['English Name']:
				elderString = elderString+" ~ "+lineDict['English Name'].strip()
			elif lineDict['English Name']:
				elderString = lineDict['English Name'].strip()
			if elderString and lineDict['Village']:
				elderString = elderString+" -- "+lineDict['Village'].strip()
			elif lineDict['Village']:
				unnamedWithVillage = True

			if elderString:		
				summariesDict[aapb]['elderTags'].append(elderString)
				elderNames.add(elderString)
				elderCat2Images[elderString].append(lineDict['Image'].strip())
			else:
				elderString = f"Yuk {unknownElderNumber:02d}"
				if unnamedWithVillage:
					elderString = elderString+" -- "+lineDict['Village'].strip()
				unknownElderNumber += 1
				elderNames.add(elderString)
				summariesDict[aapb]['elderTags'].append(elderString)
				elderCat2Images[elderString].append(lineDict['Image'].strip())


	return summariesDict, elderNames, elderCat2Images


def mixCategoriesSummaries(summaries, categories, elderCat2Images):

	# fix category tags in summaries
	for summ in summaries:
		for i, tag in enumerate(summaries[summ]["tags"]):
			if tag in categoryReplacements:
				summaries[summ]["tags"][i] = categoryReplacements[tag]

	# create video titles = Name(s), - # num
	for summ in summaries:
		newTitle = ""
		for tag in summaries[summ]["elderTags"]:
			yuk = re.sub(r' --.*','', tag.replace("~ ",""))
			if newTitle:
				newTitle += ", " + yuk
			else:
				newTitle = yuk
		newTitle += " - # " + videoNumAndDuration[summaries[summ]["videoID"]][0]
		summaries[summ]["title"] = newTitle

	# add subtitleID to summary list to be "summary":{subtitleID:[],subtitleID:[]}
	for summ in summaries:
		newSummary = {}

		# handle files without subtitle files but summary file by using negative IDs
		if not os.path.exists(os.path.join("../src/components/transcription/",summaries[summ]["videoID"]+".js")):
			print(summaries[summ]["videoID"]+".js doesn't exist; using negative IDs")
			subtitleID = -1
			for chapter in summaries[summ]["summary"]:
				newSummary[subtitleID] = chapter
				subtitleID -= 1

		else:
			# grab subtitle.js info
			subtitleTimes = {}
			with open(os.path.join("../src/components/transcription/",summaries[summ]["videoID"]+".js"), 'r') as file:
				lines = file.read()
				lines = lines.replace("export const subtitles = ","").replace("};","}")
				# print(lines)
				subtitleJSON = json.loads(lines, object_pairs_hook=OrderedDict)
				subtitleTimes = [(subtitleJSON[x]['endTime'],x) for x in subtitleJSON]

			for chapter in summaries[summ]["summary"]:
				subtitleID = ""

				# convert chapter time to seconds
				match = re.match(r'(?:(?P<hour>\d):)?(?P<minute>\d\d):(?P<second>\d\d)',chapter[0])
				if not match:
					print(f"{summ} didn't parse {chapter}")
					continue
				else:
					chapterTimeInSeconds = int(match.group('hour')) * 60 * 60 if match.group('hour') else 0
					chapterTimeInSeconds += int(match.group('minute')) * 60
					chapterTimeInSeconds += int(match.group('second'))

				# main logic
				for sub in subtitleTimes:
					if chapterTimeInSeconds > sub[0] - 1.1: 
						# print(f'{chapterTimeInSeconds}\t>=\t{sub[0]}')
						subtitleID = sub[1]
						# pass
					else:
						subtitleID = sub[1] if subtitleID != "1" else "1"
						# subtitleID = sub[1]
						break

				newSummary[subtitleID] = chapter


		summaries[summ]["summary"] = newSummary


	# print(elderCat2Images)
	for cat in categories:
		if categories[cat]["name"] in elderCat2Images.keys():
			categories[cat]["images"] = elderCat2Images[categoriesDict[cat]["name"]]

	# add base category pictures using url naming convention
	for cat in categories:
		if "." not in cat: # not a subcategory
			categories[cat]["images"].append(categories[cat]["url"].replace("-","")+".jpg")

	categories_flipped = {cat_value['name']:cat_key for cat_key, cat_value in categories.items()}
	tags_unsorted = []

	# replace tags with categoryNumbers or to 'unsorted'
	for videoNumber, summary_info in summaries.items():
		tags_new = []
		for tag in summary_info['tags']:
			if tag in categories_flipped:
				tags_new.append(categories_flipped[tag])
				categories[categories_flipped[tag]]['videoNumbers'].append(videoNumber)
			else:
				tags_new.append(tag)
				tags_unsorted.append(tag)
		eldertags_new = []
		for tag in summary_info['elderTags']:
			if tag in categories_flipped:
				eldertags_new.append(categories_flipped[tag])
				categories[categories_flipped[tag]]['videoNumbers'].append(videoNumber)
			else:
				eldertags_new.append(tag)
				tags_unsorted.append(tag)
		# summary_info['tags']
		summaries[videoNumber]['tags'] = tags_new
		summaries[videoNumber]['elderTags'] = eldertags_new

	# add videos from children to parent
	for cat_key, cat_value in categories.items():
		for i in range(1,len(cat_key.split('.'))):
			parent = cat_key.rsplit('.',i)[0]
			categories[parent]['videoNumbers'].extend(cat_value['videoNumbers']) 

	# remove duplicate videoNumbers
	for cat_key, cat_value in categories.items():
		categories[cat_key]['videoNumbers'] = list(dict.fromkeys(sorted(cat_value['videoNumbers'])))

	# remove categories without videos
	usedCategories = defaultdict(dict)
	for cat_key in categories:
		if categories[cat_key]['videoNumbers'] != []:
			usedCategories[cat_key] = categories[cat_key]
	categories = usedCategories

	categoriesUrlLookup = {cat_value['url']:cat_key for cat_key, cat_value in categories.items()}

	categories['unsorted']['tags'] = sorted(list(set(tags_unsorted)))


	return summaries, categories, categoriesUrlLookup


if __name__ == '__main__':
	categoriesFilename = '../data/categories.txt'
	summariesFilename = '../data/summaries.txt'
	elderIdentifierFilename = '../data/Elder Identification - Sheet1.tsv'
	coreyfolder = '../data/WoW-Corey'
	coreyfiles = sorted(glob.glob('../data/WoW-Corey/*.docx'))
	# coreyfiles = [x.replace('../data/WoW-Corey/','') for x in coreyfiles]


	print(f'importing {summariesFilename}')
	lonnyDict = parse_summaries(summariesFilename)
	with open('lonnysummaries.json', 'w', encoding='utf-8') as out:
		out.write(json.dumps(lonnyDict, indent=4, ensure_ascii=False))

	print(f'importing {coreyfolder}')
	coreyDict = parseCoreySummaries(coreyfiles)
	with open('coreysummaries.json', 'w', encoding='utf-8') as out:
		out.write(json.dumps(coreyDict, indent=4, ensure_ascii=False))

	summariesDict = lonnyDict | coreyDict
	# add in empty summaries
	for vid in videoNumAndDuration:
		if vid not in summariesDict:
			summariesDict[vid] = {
				"videoID":vid,
				"title":"Title Placeholder",
				"time":videoNumAndDuration[vid][1],
				"metadata":{
					"Date":"Unknown",
				},
				"summary":[],
				"elderTags":[],
				"tags":[]
			}


	# with open(os.path.join("",'summaries.js'), 'w') as out:
	# 	out.write(f"export const summaries = {json.dumps(summariesDict, indent=4)};")

	print(f'adding elder identification data')
	summariesDict, elderNames, elderCat2Images = addElderIdentification(elderIdentifierFilename, summariesDict)

	print(f'importing {categoriesFilename}')
	categoriesXML = categoriesXML(categoriesFilename, elderNames)
	categoriesDict = categoriesDictFunc(categoriesXML)


	print(f'mixing categories and summaries')
	summaries, categories, categoriesUrlLookup = mixCategoriesSummaries(summariesDict, categoriesDict, elderCat2Images)


	# TODO:
	# add where to put in summary header in subtitles (cross reference timestamps); add before "0" timestamp






	infoFolder = '../src/components/info'
	Path(infoFolder).mkdir(parents=True, exist_ok=True)

	print(f'creating categories.js, categoriesUrlLookup.js and summaries.js')
	with open(os.path.join(infoFolder,'categories.js'), 'w', encoding='utf-8') as out:
		out.write(f"export const categories = {json.dumps(categories, indent=4, ensure_ascii=False)};")

	with open(os.path.join(infoFolder,'categoriesUrlLookup.js'), 'w', encoding='utf-8') as out:
		out.write(f"export const categoriesUrlLookup = {json.dumps(categoriesUrlLookup, indent=4, ensure_ascii=False)};")

	with open(os.path.join(infoFolder,'summaries.js'), 'w', encoding='utf-8') as out:
		out.write(f"export const summaries = {json.dumps(summaries, indent=4, ensure_ascii=False)};")

	with open(os.path.join(infoFolder,'videoNum2cpb.js'), 'w', encoding='utf-8') as out:
		videoNum2cpb = {int(videoNumAndDuration[x][0]):x for x in videoNumAndDuration}
		out.write(f"export const videoNum2cpb = {json.dumps(videoNum2cpb, indent=4, ensure_ascii=False)};")


